import docx
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from string import punctuation
from nltk.stem.snowball import SnowballStemmer
import nltk
import codecs
from io import StringIO 
from docx import Document
document = Document()




#from nltk.stem import PorterStemmer

"""
For example, the Center for a New American Dream envisions "... a focus on more of what really matters, such as creating a meaningful life, contributing to community and society, valuing nature, and spending time with family and friends."
"""

#At Text we want to publish books that make a difference to people’s lives. We believe that reading should be a marvellous experience, that every book you read should somehow change your life if only by a fraction. We love the phrase ‘lost in a book’—that’s where we want our readers to be. You can’t get lost in a newspaper or a magazine or even a movie. But people get lost in books every day—on the tram, on the beach, in bed. Reading is what keeps the imagination supple and challenges preconceptions and prejudices. You read at your own speed, and the world you enter courtesy of the writer is yours and yours only, even if the person next to you on the bus is reading exactly the same book. We are committed to providing first-class services to the writers who trust us with their books. If we are to fulfil our promises to readers, our job is to do everything in our power to help our writers write the best books they can. We care passionately about the quality of our editorial work. Design and production matter to us because we want our books to be beautiful objects. And since we want our beautiful books to be read on trains and beaches, in beds and planes, or under the desk at school and work, we work up a sweat to promote and market them. And that includes selling rights on behalf of our writers to publishers all around the world. In fact we earn more money for our writers through licensing international editions than we do by turning their books into bestsellers in Australia and New Zealand.


class TextSummarizer:
    #ps = PorterStemmer()
    stemmer = SnowballStemmer("english")
    stopWords = set(stopwords.words("english")+ list(punctuation))
    text = ""
    sentences = ""
    def tokenize_sentence(self):
        words = word_tokenize(self.text)
        print(words)
        return words;

    def input_text(self):
        while True:
            # ENTER THE DOCX FILE           
            document = Document('n_thakur_petition.docx')
            #self.text = document.read()
            self.text = []
            for para in document.paragraphs:
                self.text.append(para.text)

            self.text = str(self.text)

            if(len(self.text)>10):
                break;
            else:
                print("Please input the text as length at least 10")


    def cal_freq(self,words):
        freqTable = dict()
        for word in words:
            word = word.lower()
            if word in self.stopWords:
                continue
            #word = stemmer.stem(word)

            if word in freqTable:
                freqTable[word] += 1
            else:
                freqTable[word] = 1
            return freqTable


    def compute_sentence(self,freqTable):
        self.sentences = sent_tokenize(self.text)
        sentenceValue = dict() # dict() creates the dictionary with key and it's corresponding value

        for sentence in self.sentences:

            for index, wordValue in enumerate(freqTable, start=1):

                if wordValue in sentence.lower(): # index[0] return word
                    if sentence in sentenceValue:
                        sentenceValue[sentence] += index # index return value of occurence of that word
                        #sentenceValue.update({sentence: index})
                        #print(sentenceValue)
                    else:
                        # sentenceValue[sentence] = wordValue
                        sentenceValue[sentence] = index
                        #print(sentenceValue)

        print(sentenceValue)
        return sentenceValue


    def print_summary(self,sentenceValue,average):
        self.sentences = sent_tokenize(self.text)
        sentenceValue = dict() # dict() creates the dictionary with key and it's corresponding value

        for sentence in self.sentences:

            for index, wordValue in enumerate(freqTable, start=1):

                if wordValue in sentence.lower():
                    if sentence in sentenceValue:
                        sentenceValue[sentence] += index 
                    else:
                        sentenceValue[sentence] = index
        print(sentenceValue)
        return sentenceValue



    def sumAvg(self,sentenceValue):
        sumValues = 0
        for sentence in sentenceValue:
            sumValues += sentenceValue[sentence]

        # Average value of a sentence from original text
        average = int(sumValues / len(sentenceValue))

        return average;

        summary = ''
        for sentence in self.sentences:
            if (sentence in sentenceValue) and (sentenceValue[sentence] > (1.5 * average)):
                summary += " " + sentence

        return summary

ts = TextSummarizer()
ts.input_text()

words = ts.tokenize_sentence()
freqTable = ts.cal_freq(words)
sentenceValue = ts.compute_sentence(freqTable)
avg = ts.sumAvg(sentenceValue)
summary = ts.print_summary(sentenceValue,avg)

summary = summary.replace("', '","")
summary = summary.replace(" \\t    "," ")
print(summary)
